#!/usr/bin/env python3
"""docx2md - 將 .docx 轉換為 Markdown + 提取圖片"""

import argparse
import io
import os
import re
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn

try:
    from PIL import Image

    HAS_PIL = True
except ImportError:
    HAS_PIL = False

VML_NS = "{urn:schemas-microsoft-com:vml}"


# ── 圖片提取 ──────────────────────────────────────────────


def _extract_image(rel, images_dir: Path, index: int) -> str:
    """提取單張圖片，回傳相對路徑 (images/img_001.png)。"""
    blob = rel.target_part.blob
    content_type = rel.target_part.content_type
    ext = _content_type_to_ext(content_type)

    # EMF/WMF 嘗試轉 PNG
    if ext in (".emf", ".wmf") and HAS_PIL:
        try:
            img = Image.open(io.BytesIO(blob))
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            blob = buf.getvalue()
            ext = ".png"
        except Exception:
            pass  # 轉換失敗就保留原格式

    filename = f"img_{index:03d}{ext}"
    images_dir.mkdir(parents=True, exist_ok=True)
    (images_dir / filename).write_bytes(blob)
    return f"images/{filename}"


def _content_type_to_ext(content_type: str) -> str:
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/tiff": ".tiff",
        "image/x-emf": ".emf",
        "image/x-wmf": ".wmf",
    }
    return mapping.get(content_type, ".bin")


# ── 行內格式 ──────────────────────────────────────────────


def _run_format_key(run) -> str:
    """回傳 run 的格式鍵，用於判斷相鄰 run 是否可合併。"""
    b = "b" if run.bold else ""
    i = "i" if run.italic else ""
    return f"{b}{i}"


def _wrap_text(text: str, fmt_key: str) -> str:
    """依格式鍵包裝文字。"""
    if not text:
        return ""
    # 純符號/空白不加格式標記（避免 **- ** 之類的碎片）
    if fmt_key and not any(c.isalnum() for c in text):
        return text
    if fmt_key == "bi":
        return f"***{text}***"
    if fmt_key == "b":
        return f"**{text}**"
    if fmt_key == "i":
        return f"*{text}*"
    return text


def _merge_runs(parts: list[tuple[str, str]]) -> list[str]:
    """合併相鄰同格式的 (text, fmt_key) 片段，消除碎片化標記。"""
    if not parts:
        return []
    merged = []
    cur_text, cur_fmt = parts[0]
    for text, fmt in parts[1:]:
        if fmt == cur_fmt:
            cur_text += text
        else:
            merged.append(_wrap_text(cur_text, cur_fmt))
            cur_text, cur_fmt = text, fmt
    merged.append(_wrap_text(cur_text, cur_fmt))
    return merged


def _paragraph_text(paragraph) -> str:
    """取得段落完整文字（含行內格式），不含圖片。"""
    # 先收集所有 (text, fmt_key) 片段
    raw_parts: list[tuple[str, str]] = []

    for child in paragraph._element:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "hyperlink":
            # 超連結：先 flush 合併，再加入連結
            texts = []
            for r in child.findall(qn("w:r")):
                t = r.findtext(qn("w:t"), default="")
                texts.append(t)
            link_text = "".join(texts)
            rid = child.get(qn("r:id"))
            # flush 之前的 runs
            if raw_parts:
                merged = _merge_runs(raw_parts)
                # 加入一個特殊標記讓後面能區分
                raw_parts.clear()
                link_md = ""
                if rid and rid in paragraph.part.rels:
                    url = paragraph.part.rels[rid].target_ref
                    link_md = f"[{link_text}]({url})"
                else:
                    link_md = link_text
                # 把已合併的和連結都加回去
                raw_parts.extend((t, "__literal__") for t in merged)
                raw_parts.append((link_md, "__literal__"))
            else:
                if rid and rid in paragraph.part.rels:
                    url = paragraph.part.rels[rid].target_ref
                    raw_parts.append((f"[{link_text}]({url})", "__literal__"))
                else:
                    raw_parts.append((link_text, "__literal__"))
        elif tag == "r":
            has_image = (
                child.find(".//" + qn("w:pict")) is not None
                or child.find(".//" + qn("w:object")) is not None
                or child.find(".//" + qn("w:drawing")) is not None
            )
            if not has_image:
                r = docx.text.run.Run(child, paragraph)
                # 處理 run 內的 w:br（換行符）→ 空格
                has_br = child.find(qn("w:br")) is not None
                text = r.text or ""
                if has_br and not text:
                    raw_parts.append((" ", ""))
                elif text:
                    raw_parts.append((text, _run_format_key(r)))

    # 最終合併
    if not raw_parts:
        return ""

    result = []
    buffer: list[tuple[str, str]] = []
    for text, fmt in raw_parts:
        if fmt == "__literal__":
            # flush buffer
            if buffer:
                result.extend(_merge_runs(buffer))
                buffer.clear()
            result.append(text)
        else:
            buffer.append((text, fmt))
    if buffer:
        result.extend(_merge_runs(buffer))

    return "".join(result)


# ── 圖片定位 ──────────────────────────────────────────────


def _build_image_map(document) -> dict:
    """建立 rId -> image_rel 的對照表。"""
    image_rels = {}
    for rel in document.part.rels.values():
        if "image" in rel.reltype:
            image_rels[rel.rId] = rel
    return image_rels


def _find_images_in_element(element) -> list[str]:
    """找出元素中所有圖片的 rId。"""
    rids = []
    # w:drawing > a:blip
    for blip in element.findall(".//" + qn("a:blip")):
        embed = blip.get(qn("r:embed"))
        if embed:
            rids.append(embed)
    # v:imagedata (VML)
    for imgdata in element.findall(f".//{VML_NS}imagedata"):
        rid = imgdata.get(qn("r:id"))
        if rid:
            rids.append(rid)
    return rids


def _extract_textbox_content(element) -> str | None:
    """提取元素中文字方塊的內容，若無文字方塊則回傳 None。"""
    txbx_nodes = element.findall(".//" + qn("w:txbxContent"))
    if not txbx_nodes:
        return None

    lines = []
    for txbx in txbx_nodes:
        for p in txbx.findall(qn("w:p")):
            texts = []
            for t in p.findall(".//" + qn("w:t")):
                if t.text:
                    texts.append(t.text)
            line = "".join(texts)
            lines.append(line)

    content = "\n".join(lines).strip()
    return content if content else None


# ── 表格處理 ──────────────────────────────────────────────


def _cell_text(cell) -> str:
    """取得儲存格純文字內容，多段落用空格連接，換行符替換為空格。"""
    texts = []
    for p in cell.paragraphs:
        t = _paragraph_text(p).strip()
        if t:
            texts.append(t)
    result = " ".join(texts)
    # 移除儲存格內的換行符（避免破壞 Markdown 表格格式）
    result = result.replace("\n", " ").replace("\r", " ")
    return result


def _table_to_md(table) -> str:
    """將表格轉為 Markdown 表格。單欄表格視為程式碼區塊。"""
    rows = []
    for row in table.rows:
        cells = [_cell_text(c) for c in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # 處理合併儲存格造成的重複：python-docx 會為合併格重複同一 cell
    # 去重：若連續 cell 文字相同，後面的替換為空
    cleaned_rows = []
    for row in rows:
        cleaned = []
        for i, cell in enumerate(row):
            if i > 0 and cell == row[i - 1]:
                cleaned.append("")
            else:
                cleaned.append(cell)
        cleaned_rows.append(cleaned)

    # 確保每列欄數一致
    max_cols = max(len(r) for r in cleaned_rows)
    for r in cleaned_rows:
        while len(r) < max_cols:
            r.append("")

    # 過濾全空列
    cleaned_rows = [r for r in cleaned_rows if any(c.strip() for c in r)]
    if not cleaned_rows:
        return ""

    # 單欄表格 → 程式碼區塊（常見於 Word 用表格包 SQL）
    if max_cols == 1:
        # 用原始段落文字重建，保留每段落獨立一行
        code_lines = []
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text:
                        code_lines.append(text)
        content = "\n".join(code_lines)
        return f"```sql\n{content}\n```"

    # 一般表格：轉義 pipe 字元
    for r in cleaned_rows:
        for i in range(len(r)):
            r[i] = r[i].replace("|", "\\|")

    lines = []
    # 表頭
    lines.append("| " + " | ".join(cleaned_rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    # 表身（過濾全空列）
    for row in cleaned_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


# ── 段落樣式判斷 ──────────────────────────────────────────


def _heading_level(paragraph) -> int | None:
    """回傳標題層級 (1-6)，非標題回傳 None。"""
    style_name = paragraph.style.name or ""

    # 標準 Heading 1-6
    m = re.match(r"Heading\s+(\d)", style_name)
    if m:
        return min(int(m.group(1)), 6)

    # Title → H1
    if style_name in ("Title",):
        return 1

    # Header / Heading Bar → H2
    if style_name in ("Header", "Heading Bar", "目錄標題1"):
        return 2

    return None


def _list_info(paragraph) -> tuple[int, bool] | None:
    """回傳 (縮排層級, 是否有序) 或 None。"""
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        # Bullet 樣式但沒有 numPr
        if paragraph.style.name in ("Bullet",):
            return (0, False)
        return None

    ilvl_elem = numPr.find(qn("w:ilvl"))
    level = int(ilvl_elem.get(qn("w:val"))) if ilvl_elem is not None else 0
    return (level, False)  # 簡化：全部當無序清單


# ── 主轉換 ──────────────────────────────────────────────


def _extract_metadata(body) -> dict:
    """從文件前段提取版號與日期等 metadata。"""
    meta = {"version": "", "last_updated": "", "title": ""}
    for elem in body:
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag != "p":
            continue
        texts = []
        for t in elem.findall(".//" + qn("w:t")):
            if t.text:
                texts.append(t.text)
        text = "".join(texts).strip()
        if not text:
            continue

        text_lower = text.lower()
        if text_lower.startswith("version:"):
            meta["version"] = text.split(":", 1)[1].strip()
        elif text_lower.startswith("last updated:"):
            meta["last_updated"] = text.split(":", 1)[1].strip()
        elif text_lower.startswith("creation date:"):
            if not meta["last_updated"]:
                meta["last_updated"] = text.split(":", 1)[1].strip()

        # 取文件標題（Bullet 樣式中含 SA_ 或 IM 開頭的）
        pPr = elem.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                style_val = pStyle.get(qn("w:val"), "")
                if style_val == "Bullet" and ("SA_" in text or "IM" in text):
                    meta["title"] = text

    return meta


def _find_content_start(body) -> int:
    """找出目錄（Contents/TOC）結束後第一個實質內容的 body 元素索引。"""
    toc_styles = {"21", "30", "10"}  # toc 2, toc 3, Contents heading
    found_toc = False
    last_toc_idx = 0

    for i, elem in enumerate(body):
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag != "p":
            continue

        # 檢查段落樣式
        pPr = elem.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                style_val = pStyle.get(qn("w:val"), "")
                if style_val in toc_styles:
                    found_toc = True
                    last_toc_idx = i
                    continue

        # 檢查文字是否為 "Contents"
        texts = []
        for t in elem.findall(".//" + qn("w:t")):
            if t.text:
                texts.append(t.text)
        text = "".join(texts).strip()
        if text == "Contents":
            found_toc = True
            last_toc_idx = i

    if found_toc:
        return last_toc_idx + 1
    return 0


def convert(docx_path: str, output_dir: str | None = None) -> Path:
    """轉換 docx 為 Markdown + 圖片，回傳輸出目錄。"""
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"找不到檔案: {docx_path}")

    # 決定輸出目錄
    if output_dir:
        out_dir = Path(output_dir)
    else:
        out_dir = docx_path.parent / docx_path.stem
    out_dir.mkdir(parents=True, exist_ok=True)

    images_dir = out_dir / "images"
    md_path = out_dir / f"{docx_path.stem}.md"

    document = docx.Document(str(docx_path))
    image_rels = _build_image_map(document)

    body = document.element.body

    # 提取 metadata 並找出內容起始位置
    meta = _extract_metadata(body)
    content_start = _find_content_start(body)

    # 預先提取所有圖片（只提取 content_start 之後會用到的）
    # 先收集 content_start 之後的所有 rId
    used_rids = set()
    for i, elem in enumerate(body):
        if i < content_start:
            continue
        used_rids.update(_find_images_in_element(elem))

    extracted_images: dict[str, str] = {}
    img_counter = 0
    for rid, rel in image_rels.items():
        if rid in used_rids:
            img_counter += 1
            extracted_images[rid] = _extract_image(rel, images_dir, img_counter)

    # 建立 metadata 標頭
    from datetime import datetime

    md_lines: list[str] = []
    title = meta["title"] or docx_path.stem
    md_lines.append(f"# {title}")
    md_lines.append("")
    meta_parts = []
    if meta["version"]:
        meta_parts.append(f"版本：{meta['version']}")
    if meta["last_updated"]:
        meta_parts.append(f"文件日期：{meta['last_updated']}")
    meta_parts.append(f"轉換時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    md_lines.append(" | ".join(meta_parts))
    md_lines.append("")
    md_lines.append("---")
    md_lines.append("")

    # 從 content_start 開始遍歷 body 元素
    for i, element in enumerate(body):
        if i < content_start:
            continue
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            _process_paragraph(element, document, extracted_images, md_lines)
        elif tag == "tbl":
            _process_table(element, document, extracted_images, md_lines)

    # 寫入 Markdown
    content = "\n".join(md_lines).strip() + "\n"
    content = re.sub(r"\n{4,}", "\n\n\n", content)
    md_path.write_text(content, encoding="utf-8")

    return out_dir


def _process_paragraph(element, document, extracted_images, md_lines):
    """處理單一段落元素。"""
    paragraph = docx.text.paragraph.Paragraph(element, document)

    # 先處理段落中的文字方塊
    txbx_content = _extract_textbox_content(element)

    # 再處理段落中的圖片
    img_rids = _find_images_in_element(element)
    if img_rids:
        for rid in img_rids:
            if rid in extracted_images:
                md_lines.append(f"![{extracted_images[rid]}]({extracted_images[rid]})")
                md_lines.append("")
        # 如果段落只有圖片沒有文字，輸出文字方塊後返回
        text = _paragraph_text(paragraph).strip()
        if not text:
            if txbx_content:
                md_lines.append(f"```\n{txbx_content}\n```")
                md_lines.append("")
            return

    text = _paragraph_text(paragraph).strip()

    # 空段落
    if not text:
        md_lines.append("")
        return

    # 標題
    level = _heading_level(paragraph)
    if level is not None:
        md_lines.append("")
        md_lines.append(f"{'#' * level} {text}")
        md_lines.append("")
        return

    # 清單
    list_info = _list_info(paragraph)
    if list_info is not None:
        indent_level, is_ordered = list_info
        indent = "  " * indent_level
        prefix = "1. " if is_ordered else "- "
        md_lines.append(f"{indent}{prefix}{text}")
        return

    # 一般段落
    md_lines.append(text)
    md_lines.append("")

    # 段落後附加文字方塊內容
    if txbx_content:
        md_lines.append(f"```\n{txbx_content}\n```")
        md_lines.append("")


def _process_table(element, document, extracted_images, md_lines):
    """處理單一表格元素。"""
    table = docx.table.Table(element, document)

    # 檢查表格中是否有圖片
    img_rids = _find_images_in_element(element)
    for rid in img_rids:
        if rid in extracted_images:
            md_lines.append(f"![{extracted_images[rid]}]({extracted_images[rid]})")
            md_lines.append("")

    md_table = _table_to_md(table)
    if md_table:
        md_lines.append("")
        md_lines.append(md_table)
        md_lines.append("")


# ── CLI ───────────────────────────────────────────────────


def main():
    parser = argparse.ArgumentParser(
        prog="docx2md",
        description="將 .docx 轉換為 Markdown + 提取圖片",
    )
    parser.add_argument("input", help="輸入的 .docx 檔案路徑")
    parser.add_argument("-o", "--output", help="輸出目錄（預設：與輸入檔同名的資料夾）")

    args = parser.parse_args()

    if not args.input.lower().endswith(".docx"):
        print("錯誤：僅支援 .docx 格式", file=sys.stderr)
        sys.exit(1)

    try:
        out_dir = convert(args.input, args.output)
        print(f"轉換完成：{out_dir}")
        # 列出產生的檔案
        for f in sorted(out_dir.rglob("*")):
            if f.is_file():
                size = f.stat().st_size
                rel = f.relative_to(out_dir)
                print(f"  {rel}  ({size:,} bytes)")
    except FileNotFoundError as e:
        print(f"錯誤：{e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"轉換失敗：{e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
